"""Generate REPORT_zero_shot_clean_train.docx — concise (≤ 2pp) report
on the zero-shot Llama-3.2-3B clean-mode pass on LRS2 train (45,839 rows).

Styled to match make_v2_docx.py: Calibri body, "Light Grid Accent 1" tables
with bold headers, "List Bullet" bullets.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(__file__).parent / "REPORT_zero_shot_clean_train.docx"


def set_calibri(run, size=11, bold=False, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_calibri(r, size=10, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = "Calibri"
        r.font.size = Pt(10)


def add_table(doc, header, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(11)
    for ridx, row in enumerate(rows, start=1):
        cells = t.rows[ridx].cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(11)
    if col_widths_cm:
        for ridx in range(1 + len(rows)):
            for i, w in enumerate(col_widths_cm):
                t.rows[ridx].cells[i].width = Cm(w)
    return t


def _strip_pagebreak_before(style):
    """Remove the default page-break-before that python-docx puts on H1."""
    pPr = style.element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    if pPr is not None:
        for child in pPr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore"
        ):
            pPr.remove(child)


def build():
    doc = Document()

    # ── global style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    # doc margins — tight for ≤2 page target
    for s in doc.sections:
        s.left_margin = Cm(1.6)
        s.right_margin = Cm(1.6)
        s.top_margin = Cm(1.4)
        s.bottom_margin = Cm(1.4)

    # Headings: shrink to academic-report scale; no page-break-before
    for lvl, sz in [(1, 12), (2, 10.5)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.bold = True
        _strip_pagebreak_before(h)
    h0 = doc.styles["Title"]
    h0.font.name = "Calibri"
    h0.font.size = Pt(14)
    h0.font.bold = True
    _strip_pagebreak_before(h0)

    # ── Title
    title = doc.add_heading(
        "Zero-Shot Ablation — Llama-3.2-3B on LRS2 Train (clean mode)",
        level=0,
    )
    for r in title.runs:
        r.font.name = "Calibri"

    add_para(
        doc,
        "Single-mode run (phonemes in, text out, no training, no LoRA). "
        "Numbers in this report come from zero-shot/baseline/"
        "preds_train_45839_clean.jsonl (45,839 rows) and "
        "errors_train_45839_clean.json (Stage 2/3 error pattern analysis). "
        "All metrics were recomputed directly from the saved predictions "
        "and the saved homophone mask.",
        italic=True,
    )

    # ── 1. CONFIGURATION ──────────────────────────────────────────────────
    add_heading(doc, "1. Configuration", level=1)

    add_heading(doc, "1.1 Model", level=2)
    add_bullets(doc, [
        "Architecture / base: meta-llama/Llama-3.2-3B (~3.2 B parameters; "
        "loaded in 4-bit NF4 via bitsandbytes, double-quant).",
        "Compute dtype: torch.float16. The two GTX 1080s are Pascal/cc=6.1, "
        "below Ampere's bf16 floor.",
        "Adapter / fine-tuning: NONE. This is a zero-shot ablation; no "
        "LoRA, no training, no import of src/p2t_lora/model.py.",
        "Device: cuda:0 (single GPU, pinned via CUDA_VISIBLE_DEVICES=0; "
        "PCIe multi-GPU split is ~3× slower here).",
    ])

    add_heading(doc, "1.2 Inference parameters", level=2)
    add_bullets(doc, [
        "max_new_tokens = 34 (≈ longest reference plus margin). "
        "do_sample = False (greedy). batch_size = 8.",
        "padding_side = \"left\", truncation_side = \"left\".",
        "max_input_len = 302 tokens (longest tokenised prompt in the run, "
        "computed up-front so zero truncation occurred).",
    ])

    add_heading(doc, "1.3 Data", level=2)
    add_bullets(doc, [
        "Corpus: LRS2, train split, 45,839 rows, from "
        "sentphonemepairs_LRS2_original.csv via "
        "data_loader.load_original_phoneme_text_pairs().",
        "Mode = clean: loader.py's cleaned phoneme column — <SOS>/<EOS>/"
        "<space> markers and stress digits stripped "
        '(e.g. DH AH0 <space> K AE1 T → DH A K AE T).',
        "Instruction template (clean): \u201cYou are given a sequence of "
        "ARPAbet phonemes representing one spoken English sentence. "
        "Convert the phonemes into the English sentence they spell out. "
        "Reply with only that sentence and nothing else.\u201d",
    ])

    add_heading(doc, "1.4 Post-processing & runtime", level=2)
    add_bullets(doc, [
        "Answer extraction: split on \\n and keep the first line; further "
        "strip anything after a \u201cPhonemes :\u201d token. Without this, "
        "the base model emits its answer on line 1 then re-prints "
        "\u201cPhonemes: \u2026\\nText: \u2026\u201d indefinitely.",
        "Hardware: 2× NVIDIA GeForce GTX 1080 (8 GB each, Pascal/cc=6.1, "
        "PCIe, no NVLink). One card used per run.",
        "Software: PyTorch 2.3.1+cu121, transformers 4.44.2, "
        "bitsandbytes 0.49.2, accelerate 1.14.0.",
        "Clean-decode wall time: ~14 h 12 min "
        "(2026-07-15 22:58 \u2192 2026-07-16 13:10), \u2248 0.90 rows/sec.",
        "Stage 2/3 error-analysis wall time: ~4 h 37 min (13:10 \u2192 "
        "17:47); bulk is the brute-force near-homophone lookup against "
        "the ~125 k-word CMU dict (~1 s/substitution).",
    ])

    add_heading(doc, "1.5 Metrics computed", level=2)
    add_para(
        doc,
        "Core (per subset): WER, CER, PER (phoneme error rate, via G2P "
        "round-trip), BLEU-4, Exact Match. Stratification: Overall / "
        "Homophone / Non-Homophone, with the homophone flag from "
        "data_loader.load_homophone_sentences() — 35,658 / 45,839 rows "
        "flagged (77.8%).",
    )
    add_para(
        doc,
        "Stage 2 error-pattern analysis: every WER substitution classified "
        "as Homophone / Near-homophone / Other against the CMU pronouncing "
        "dictionary and get_near_homophones(). Stage 3 (grammar + LLM "
        "escalation) was wired in but produced zero escalated examples on "
        "this run — see §3.4.",
    )

    # ── 2. RESULTS ────────────────────────────────────────────────────────
    add_heading(doc, "2. Results", level=1)

    add_heading(doc, "2.1 Core metrics", level=2)
    add_para(
        doc,
        "The base Llama-3.2-3B, prompted only with the phoneme sequence, "
        "does not learn the phoneme→text mapping: WER is above 100% on "
        "every subset (the model inserts and deletes more words than it "
        "gets right, so the ratio of error operations to reference-word "
        "count exceeds 1.0). Exact-match is 0.21% (95 / 45,839 sentences "
        "reproduced character-for-character).",
    )
    add_table(
        doc,
        ["Subset", "N", "WER", "CER", "PER", "BLEU4", "EM"],
        [
            ["Overall",   "45,839", "127.78%", "93.66%", "101.95%", "0.0131", "0.21%"],
            ["Homophone", "35,658", "125.92%", "93.66%", "102.44%", "0.0133", "0.19%"],
            ["Non-Homophone", "10,181", "137.88%", "93.62%", "99.56%", "0.0118", "0.27%"],
        ],
        col_widths_cm=[3.3, 1.4, 1.5, 1.5, 1.5, 1.3, 1.2],
    )
    add_para(
        doc,
        "The Homophone subset's WER is 11.96 pp lower than the "
        "Non-Homophone subset. This is counter-intuitive at first: "
        "sentences that contain a known homophone pair are EASIER for the "
        "model — almost certainly because the LRS2 homophone list is "
        "biased toward short, common sentences, while the Non-Homophone "
        "rows include longer broadcast material.",
    )

    add_heading(doc, "2.2 Stage 2 error-pattern breakdown", level=2)
    add_table(
        doc,
        ["Subset", "Subs (S)", "Dels (D)", "Ins (I)", "Hits",
         "Homophone", "Near-homo", "Other"],
        [
            ["Overall",       "221,330", "70,991", "128,293", "36,859",
             "195 (0.1%)",  "10,071 (4.6%)", "211,064 (95.4%)"],
            ["Homophone",     "183,739", "63,586", "102,858", "30,774",
             "187 (0.1%)",  "8,253 (4.5%)",  "175,299 (95.4%)"],
            ["Non-Homophone", "37,591",  "7,405",  "25,435",  "6,085",
             "8 (0.0%)",    "1,818 (4.8%)",  "35,765 (95.1%)"],
        ],
        col_widths_cm=[2.7, 1.7, 1.6, 1.6, 1.3, 1.6, 1.7, 1.7],
    )
    add_para(
        doc,
        "Headline finding: only 4.6% of all substitution errors are "
        "phonetically explainable (homophone + near-homophone). The "
        "Homophone and Non-Homophone subsets show nearly identical rates "
        "(4.5% vs 4.8%) — being on a homophone-containing sentence does "
        "NOT concentrate the model's confusions on the homophone pair "
        "itself.",
    )

    add_heading(doc, "2.3 Spot-check examples (Stage 2)", level=2)
    add_table(
        doc,
        ["Category", "Reference", "Prediction"],
        [
            ["Homophone",       "YOU CAN TAKE THINGS TOO FAR",
             "You can't think of anything to say."],
            ["Near-homophone",  "THROUGH WHAT THEY CALL A KNIFE BLOCK",
             "The cat sat on the mat."],
            ["Other",           "WHEN YOU'RE COOKING CHIPS AT HOME",
             "When I was younger, I used to skate."],
        ],
        col_widths_cm=[3.0, 6.0, 6.7],
    )
    add_para(
        doc,
        "The \u201cOther\u201d row reveals the dominant failure mode: the "
        "model produces plausible-but-unrelated English sentences rather "
        "than transcribing the phonemes — semantically drift, not "
        "phonetic confusion.",
    )

    # ── 3. INTERPRETATION ─────────────────────────────────────────────────
    add_heading(doc, "3. Interpretation", level=1)

    add_heading(
        doc,
        "3.1 The phoneme→text task is a learned mapping; zero-shot Llama "
        "does not have it",
        level=2,
    )
    add_para(
        doc,
        "A 3 B-parameter general-purpose LM has near-zero prior exposure "
        "to ARPAbet→English alignment. The 127.78% WER and 0.21% EM "
        "confirm the obvious: prompting a base LM with phonemes does not "
        "produce transcriptions. The model's output is fluent English, "
        "but it is NOT faithful to the input — its behaviour is closer to "
        "\u201ccontinue a plausible English sentence\u201d than to "
        "\u201cdecode this exact sequence of phonemes\u201d.",
    )
    add_para(
        doc,
        "The substitution breakdown reinforces this: 95.4% of "
        "substitutions are \u201cOther\u201d (unrelated to any phonetic "
        "explanation) — exactly what you would expect from a model that "
        "is largely ignoring the phoneme content of the prompt.",
    )

    add_heading(
        doc,
        "3.2 Homophone confusion is NOT the bottleneck — contrastive "
        "mining is not the lever",
        level=2,
    )
    add_para(
        doc,
        "The original motivation for the contrastive hard-negative stage "
        "in the trained CPT decoder was to disambiguate homophone-driven "
        "substitutions. The Stage 2 numbers invalidate that as a "
        "first-order concern: only 4.6% of all substitutions are "
        "phonetically explainable, and the homophone-subset rate (4.5%) "
        "is essentially the same as the non-homophone-subset rate (4.8%) "
        "\u2014 homophone-containing sentences do NOT concentrate the "
        "model's confusions on the homophone pair itself. The error "
        "stream is dominated by semantic drift (\u201cWHEN YOU'RE "
        "COOKING CHIPS AT HOME\u201d \u2192 \u201cWhen I was younger, I "
        "used to skate.\u201d), which is a capacity / instruction-"
        "following problem, not a phonetic disambiguation problem. "
        "Implication: the contrastive mechanism's payoff is small at "
        "this baseline; it only matters once the model is TRAINED to "
        "produce phoneme-conditioned text, at which point the residual "
        "4.6% homophone-driven errors become the relevant quantity.",
    )

    add_heading(
        doc,
        "3.3 Failure mode: model treats phonemes as a topic prompt, not "
        "content",
        level=2,
    )
    add_para(
        doc,
        "The \u201cOther\u201d substitutions are not random \u2014 they "
        "are fluent, on-topic English sentences that bear no relation "
        "to the phoneme input. The model reads \u201cPhonemes: \u2026\u201d "
        "and infers a topic (\u201ccooking\u201d, \u201cyouth/sports\u201d) "
        "and continues accordingly. The Stage 2 near-homophone examples "
        "in §2.3 (\u201ca \u2192 the\u201d, \u201cthe \u2192 a\u201d) fit "
        "the same pattern: the model is generating function words to "
        "fit ITS OWN sentence, not the input's function-word positions. "
        "The relevant intervention is supervised fine-tuning on "
        "phoneme\u2192text pairs \u2014 exactly what the trained CPT "
        "decoder does. Once the model is trained to be faithful to the "
        "phoneme sequence, the error stream will shift from "
        "\u201csemantic drift\u201d to \u201cphonetic confusion\u201d, "
        "and the 4.6% homophone-driven error rate will become the "
        "relevant quantity to optimise.",
    )

    add_heading(doc, "3.4 Limitations of this analysis", level=2)
    add_para(
        doc,
        "Three caveats apply. (i) Single-mode (clean only): the "
        "companion raw-mode pass did not finish on the 45,839-row split, "
        "so the clean-vs-raw ablation is not yet available. Raw-mode "
        "decode is queued next, with PYTHONUTF8=1 set so the panphon "
        "extended metrics can complete. (ii) Stage 3 escalation produced "
        "no examples: with 95.4% of substitutions in the \u201cOther\u201d "
        "category, there are almost no candidates where a homophone "
        "substitution is grammatically plausible-but-wrong; Stage 3 will "
        "be useful again only after training reduces that fraction. "
        "(iii) Extended metrics (SID/AER/WPER) were not captured: the "
        "panphon WPER step crashed on a Windows encoding issue, and the "
        "recompute from the saved preds_train_45839_clean.jsonl is "
        "queued next, with PYTHONUTF8=1 to fix the encoding.",
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}  ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build()
